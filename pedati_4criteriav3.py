import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO

# --- 1. CONFIGURATION ---
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

@st.cache_resource
def find_working_model():
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except:
        return "models/gemini-1.5-flash"
    return "models/gemini-1.5-flash"

selected_model_name = find_working_model()
model = genai.GenerativeModel(selected_model_name)

# --- 2. AI LOGIC (INTEGRATED CRITERIA WITH COMBINED PEDATI FLOW) ---
def generate_advanced_plan(topic, syllabus, extra_context):
    prompt = f"""
    Topic: {topic}. Syllabus Code: {syllabus}. Context: {extra_context}.
    Generate a professional lesson plan in English.
    
    Use the following EXACT markers for the document structure:
    
    SECTION: TOPIC
    {topic}
    
    SECTION: LESSON OBJECTIVES
    [4 points]
    
    SECTION: LESSON OUTCOMES
    [4 points]
    
    SECTION: SUCCESS CRITERIA
    [4 points]
    
    SECTION: PREREQUISITE
    [1 point]
    
    SECTION: KEYWORDS
    [6 items]
    
    SECTION: HOTS
    [4 main domains from Bloom's Taxonomy]
    
    SECTION: DIGITAL CITIZENSHIP
    [4 points on ethical tech use/Chromebooks/Canva/YouTube]

    SECTION: OPENING LESSON CONTENT
    [Hook activity and transition plan]

    SECTION: DIFFERENTIATION STRATEGIES (GREEN)
    - HA (Higher Achiever): [1 challenging activity]

    SECTION: DIFFERENTIATION STRATEGIES (YELLOW)
    - MA (Medium Achiever): [1 core activity]

    SECTION: DIFFERENTIATION STRATEGIES (RED)
    - LA (Lower Achiever): [1 scaffolded activity]

    SECTION: BLENDED LEARNING Activity ONE (15 MINS)
    - Activity 1: [Descriptions]
    - Teacher Preparation: [Step-by-step before lesson]
    - Objectives: [3 points]
    - Student Tasks: [Step-by-step details]

    SECTION: BLENDED LEARNING Activity TWO (15 MINS)
    - Activity 2: [Descriptions]
    - Teacher Preparation: [Step-by-step before lesson]
    - Objectives: [3 points]
    - Student Tasks: [Step-by-step details]

    SECTION: PEDATI FLOW GRID
    [Generate content for the 4 pedagogical blocks exactly using the layout below. Keep descriptions concise and practical.]
    BLOCK_START: P: PREPARATION (LEARN)
    LECTURER: [Actionable steps aligned with the topic]
    STUDENTS: [Actionable tasks/chromebook work aligned with the topic]
    BLOCK_END
    
    BLOCK_START: E: ENGAGE (EXPLORE)
    LECTURER: [Actionable steps aligned with the topic]
    STUDENTS: [Actionable tasks/chromebook work aligned with the topic]
    BLOCK_END

    BLOCK_START: D.A: DELIVER AND APPLY
    LECTURER: [Actionable steps aligned with the topic]
    STUDENTS: [Actionable tasks/chromebook work aligned with the topic]
    BLOCK_END

    BLOCK_START: T.I: TEST AND EVALUATE
    LECTURER: [Actionable steps aligned with the topic]
    STUDENTS: [Actionable tasks/chromebook work aligned with the topic]
    BLOCK_END
    
    SECTION: PLENARY (EXIT TICKET)
    [2-3 minute closing activity]

    SECTION: HOMEWORK
    [Task assigned based on topic]

    SECTION: SUGGESTED WAY FORWARD TASK
    [Hook activity and transition plan for next day lesson]
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"System Error: {str(e)}"

# --- 3. WORD EXPORT LOGIC (WITH NEW GRID PARSER) ---
def create_word_export(topic, syllabus, text):
    doc = Document()
    doc.add_heading(f'PTES Lesson Plan: {topic}', 0)

    # Admin Header Table
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Week No:", "Date:"], ["Class Size:", "Day:"], ["Venue:", "Duration:"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
    doc.add_paragraph()

    # Split output into major structural sections
    sections = text.split('SECTION:')
 
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        title = lines[0].strip().replace("**", "")
        body_content = "\n".join(lines[1:]).strip()

        # Check if this section contains the specialized visual tables
        if "PEDATI FLOW GRID" in title.upper():
            doc.add_heading("P.E.D.A.T.I Flow Breakdown", level=1)
            
            # Parse individual blocks
            blocks = body_content.split("BLOCK_START:")
            for block in blocks:
                if not block.strip(): continue
                block_data = block.split("BLOCK_END")[0].strip().split('\n')
                
                heading_title = block_data[0].strip().replace("**", "")
                lecturer_text = ""
                students_text = ""
                
                for line in block_data:
                    if line.upper().startswith("LECTURER:"):
                        lecturer_text = line.split(":", 1)[1].strip().replace("**", "")
                    elif line.upper().startswith("STUDENTS:"):
                        students_text = line.split(":", 1)[1].strip().replace("**", "")
                
                # Render Table Heading
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                run = p.add_run(heading_title)
                run.bold = True
                run.font.size = Pt(12)
                
                # Draw the customized layout table
                table = doc.add_table(rows=2, cols=2)
                table.style = 'Table Grid'
                
                # Setup Column Widths
                for row in table.rows:
                    row.cells[0].width = Inches(3.25)
                    row.cells[1].width = Inches(3.25)
                
                # Header Row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Lecturer"
                hdr_cells[1].text = "Students"
                hdr_cells[0].paragraphs[0].runs[0].font.italic = True
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[1].paragraphs[0].runs[0].font.italic = True
                hdr_cells[1].paragraphs[0].runs[0].font.bold = True
                
                # Data Row
                data_cells = table.rows[1].cells
                data_cells[0].text = lecturer_text
                data_cells[1].text = students_text
        else:
            # Standard single box rendering engine with automatic asterisk cleaning
            content = body_content.replace("**", "") 
            doc.add_heading(title.title(), level=1)
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            table.cell(0, 0).text = content
            doc.add_paragraph()
     
    # HOD Approval Table Block
    doc.add_page_break()
    doc.add_heading("HOD Approval & Remarks", level=1)
    hod_table = doc.add_table(rows=2, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Remarks:"
    hod_table.rows[1].height = Pt(50)
    hod_table.cell(1, 0).text = "Date:"; hod_table.cell(1, 1).text = "Signature:"

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 4. GUI ---
st.set_page_config(page_title="Advanced Lesson Planner", layout="wide")

st.title("🎓 PTES Universal Lesson Planner (With PEDATI Flow Grid)")
st.info("Type in the lesson topic, the subject's syllabus code and the extra information like canva, youtube, infographic")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Lesson Topic:")
with c2: u_syllabus = st.text_input("Syllabus Code:")
u_extra = st.text_area("Extra Context (Optional):")

if st.button("🚀 GENERATE COMPLETE LESSON PLAN"):
    if u_topic and u_syllabus:
        with st.spinner("AI is integrating all criteria into your plan..."):
            result = generate_advanced_plan(u_topic, u_syllabus, u_extra)
            st.session_state['adv_plan_out'] = result
    else:
        st.warning("Please fill in the Topic and Syllabus.")

if 'adv_plan_out' in st.session_state:
    st.divider()
    st.subheader("AI Draft Preview")
    st.text_area("Content", st.session_state['adv_plan_out'], height=400)
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['adv_plan_out'])
    st.download_button("📥 Download to Word version (.docx)", doc_file, f"Universal_LP_{u_topic}.docx")

st.markdown("---")
st.caption("Lesson planner 4.0 | Developer: Hjh Nurul Haziqah Hj Nordin | © 2026 PTES Innovation")
