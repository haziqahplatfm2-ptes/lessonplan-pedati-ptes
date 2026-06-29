# This is the updated python script with custom document geometries, formatting rules, and keyword tables.
import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from io import BytesIO

# --- 1. CONFIGURATION ---
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

@st.cache_resource
def find_working_model():
    """Universal loader to find a working model and avoid 404 version errors."""
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except:
        return "models/gemini-1.5-flash"
    return "models/gemini-1.5-flash"

selected_model_name = find_working_model()
model = genai.GenerativeModel(selected_model_name)


def generate_pedati_plan(topic, syllabus, extra_context):
    # Rule 2, 4, and 5 are strictly enforced via the prompt rules below
    prompt = f"""
    Topic: {topic}. Syllabus Code: {syllabus}. Context: {extra_context}.
    Generate a professional lesson plan in English. 
    NO Malay terms. Use these exact stage names:
    P [Prior Knowledge], E [Engage], D [Develop], A [Apply], T [Test], I [Improve].

    CRITICAL RULES FOR CONTENT FORMATTING:
    1. DO NOT use double asterisks (**) anywhere in your response. 
    2. DO NOT use bullet points (e.g., -, *, •) under any circumstances. If a section requires multiple items or a list, you MUST use numbers (1, 2, 3...) exclusively.
    3. All heading markers below must remain in absolute CAPITAL LETTERS.

    Structure with these markers for boxing:
    SECTION: LESSON OBJECTIVES
    [Provide exactly 4 numbered points using 1., 2., 3., 4.]
    
    SECTION: LESSON OUTCOMES
    [Provide exactly 4 numbered points using 1., 2., 3., 4.]
    
    SECTION: SUCCESS CRITERIA
    [Provide exactly 4 numbered points using 1., 2., 3., 4.]
    
    SECTION: PREREQUISITE
    [Provide 1 statement]
    
    SECTION: KEYWORDS
    [Provide 6 items separated by commas only. Do not make a list.]
    
    SECTION: HOTS
    [Provide exactly 4 numbered items indicating domains in Bloom's taxonomy]
    
    SECTION: DIGITAL CITIZENSHIP
    [Provide exactly 4 numbered points using 1., 2., 3., 4. on the use of online resources like youtube channel or canva application or use of chromebooks or use of digital devices]

    SECTION: PEDATI STAGES
    STAGE: P [Prior Knowledge] | SB: [Activity] | CB: [Activity]
    STAGE: E [Engage] | SB: [Activity] | CB: [Activity]
    STAGE: D [Develop] | SB: [Activity] | CB: [Activity]
    STAGE: A [Apply] | SB: [Activity] | CB: [Activity]
    STAGE: T [Test] | SB: [Activity] | CB: [Activity]
    STAGE: I [Improve] | SB: [Activity] | CB: [Activity]
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"System Error: {str(e)}"


def add_page_number(run):
    """Helper function to inject dynamic Word field codes for top page numbers."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def create_word_export(topic, syllabus, text):
    doc = Document()
    
    # Rule 1: Set Letter Paper Size and 0.5-inch Margins on all 4 corners
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Rule 7: Set Page Numbering in Header, Top and Centered
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_p.add_run()
        header_run.font.name = 'Arial'
        header_run.font.size = Pt(10)
        add_page_number(header_run)

    # Rule 2 & 8: Main Title in CAPITAL LETTERS with Font Size 14
    main_title = f'LESSON PLAN: {topic} ({syllabus})'.upper()
    title_p = doc.add_heading(level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(main_title)
    title_run.font.size = Pt(14)
    title_run.bold = True
    
    # Rule 3: Single spacing configuration for document body styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12) # Rule 8: Rest of the text using font size 12
    style.paragraph_format.line_spacing = 1.0  # Single line spacing
    style.paragraph_format.space_after = Pt(0)  # Remove extra spacing after paragraphs
    style.paragraph_format.space_before = Pt(0)

    # 1. Admin Header Table
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Week No :", "Date:"], ["No. of Students:", "Day:"], ["Venue / Lab No:", "Duration (mins):"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
    
    # Re-apply font rules and single spacing to admin header table text
    for row in admin_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.size = Pt(12)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.line_spacing = 1.0

    # 2. Resources Table
    # Rule 2: Title in CAPITAL LETTERS with size 14
    r_heading = doc.add_paragraph()
    r_heading.paragraph_format.line_spacing = 1.0
    r_run = r_heading.add_run("RESOURCES & MATERIALS")
    r_run.bold = True
    r_run.font.size = Pt(14)
    
    res_table = doc.add_table(rows=1, cols=1)
    res_table.style = 'Table Grid'
    res_table.cell(0, 0).text = "Smart board, Chromebook, Writing table, Projector, Screen share with laptop"
    res_table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = 1.0

    # 3. Content Parsing & Table Boxing
    sections = text.split('SECTION:')
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        
        # Rule 2 & 4: Capitalizing Heading Titles and dropping any residual asterisks
        title = lines[0].strip().replace("**", "").upper()
        content_lines = lines[1:]
        
        doc_heading = doc.add_paragraph()
        doc_heading.paragraph_format.line_spacing = 1.0
        h_run = doc_heading.add_run(title)
        h_run.bold = True
        h_run.font.size = Pt(14)  # Rule 8: Section titles font size 14

        # Rule 6: Handle Keyword grid restructuring
        if "KEYWORDS" in title:
            # Join content lines and split into clean individual keyword tokens
            raw_keywords_text = " ".join([l.strip() for l in content_lines if l.strip()])
            keyword_items = [kw.strip() for kw in raw_keywords_text.split(",") if kw.strip()]
            
            # Create a clean grid table for keywords (e.g., 2 rows, 3 columns matrix)
            kw_table = doc.add_table(rows=2, cols=3)
            kw_table.style = 'Table Grid'
            
            idx = 0
            for r in range(2):
                for c in range(3):
                    if idx < len(keyword_items):
                        cell = kw_table.cell(r, c)
                        cell.text = keyword_items[idx]
                        # Aligned to the center
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1.0
                        if p.runs:
                            p.runs[0].font.size = Pt(12)
                        idx += 1
            doc.add_paragraph().paragraph_format.line_spacing = 1.0
            
        elif "|" in section and "PEDATI" in title:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Stage (PEDATI)', 'Activity One ', 'Activity Two '
            
            # Formatting table headings
            for cell in hdr:
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                if p.runs:
                    p.runs[0].font.size = Pt(12)
                    p.runs[0].font.bold = True

            for line in content_lines:
                if "|" in line:
                    p_split = line.split("|")
                    row = table.add_row().cells
                    # Rule 4: Explicit clean text strings inside table content
                    row[0].text = p_split[0].split(":")[-1].strip().replace("**", "")
                    row[1].text = p_split[1].split(":")[-1].strip().replace("**", "")
                    row[2].text = p_split[2].split(":")[-1].strip().replace("**", "")
                    
                    for cell in row:
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.0
                        if p.runs:
                            p.runs[0].font.size = Pt(12)
            doc.add_paragraph().paragraph_format.line_spacing = 1.0
        else:
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            
            # Rule 4: Clean all standard block text assignments of raw markers
            cleaned_content = "\n".join([l.strip() for l in content_lines if l.strip()]).replace("**", "")
            table.cell(0, 0).text = cleaned_content
            
            p = table.cell(0, 0).paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            if p.runs:
                p.runs[0].font.size = Pt(12)
            doc.add_paragraph().paragraph_format.line_spacing = 1.0

    # 4. HOD Approval Page
    doc.add_page_break()
    
    # Rule 2 & 8: Section Title in CAPITAL LETTERS, size 14
    hod_heading = doc.add_paragraph()
    hod_heading.paragraph_format.line_spacing = 1.0
    hod_run = hod_heading.add_run("HOD APPROVAL & REMARKS")
    hod_run.bold = True
    hod_run.font.size = Pt(14)
    
    hod_table = doc.add_table(rows=3, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Remark"
    hod_table.cell(0, 1).text = "Signature / Stamp"
    hod_table.rows[1].height = Pt(60)
    hod_table.cell(2, 0).text = "Date:"
    hod_table.cell(2, 1).text = "Name:"
    
    for row in hod_table.rows:
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            if p.runs:
                p.runs[0].font.size = Pt(12)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 5. GUI SECTION ---
st.set_page_config(page_title="PEDATI Master Planner", layout="wide")

with st.sidebar:
    st.title("📖 User Guide")
    st.info("How to use this portal:")
    st.markdown("""
    ### 1. Fill Details
    Enter your **Lesson Topic** and **Syllabus Code**. Use the **Context** box for specific requirements like "Group Work" or "Case Study".
    
    ### 2. Generate
    Click **🚀 GENERATE** and wait for the AI to draft your plan.
    
    ### 3. Review & Save
    Check the **AI Preview**. If it looks good, click **📥 Download Word** to get your document.
    """)
    st.markdown("---")
    st.caption("App Version 3.0 | Custom Layout Engine")

st.title("🎓 PEDATI Lesson Plan Generator")
st.info(f"System connected via: {selected_model_name}")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Lesson Topic:")
with c2: u_syllabus = st.text_input("Syllabus Code:")
u_extra = st.text_area("Specific Context/Keywords (Optional):")

if st.button("🚀 GENERATE PEDATI LESSON PLAN"):
    if u_topic and u_syllabus:
        with st.spinner("AI is building your structured PEDATI plan..."):
            result = generate_pedati_plan(u_topic, u_syllabus, u_extra)
            # Instantly clean double asterisks from the preview window state
            st.session_state['pedati_out'] = result.replace("**", "")
    else:
        st.warning("Please fill in the Topic and Syllabus.")

if 'pedati_out' in st.session_state:
    st.divider()
    st.text_area("AI Preview (Clean No Asterisk format)", st.session_state['pedati_out'], height=300)
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['pedati_out'])
    st.download_button("📥 Download Word (.docx)", doc_file, f"PEDATI_{u_topic}.docx")

st.markdown("---") 
st.markdown(
    """
    <div style='text-align: center; color: grey; font-size: 0.8em;'>
        <p><b>Smart PEDATI Lesson Plan AI-Generator v3.0</b></p>
        <p>Developed & Conceptualized by: <b>Hajah Nurul Haziqah @ Hjh Hartini Hj Nordin</b></p>
        <p>© 2026 PTES Academic Innovation Computer Science</p>
    </div>
    """,
    unsafe_allow_title=True, unsafe_allow_html=True
)
