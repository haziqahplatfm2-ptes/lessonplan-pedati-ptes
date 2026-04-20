import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from io import BytesIO

# --- 1. CONFIGURATION ---
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

@st.cache_resource
def find_working_model():
    """Universal loader to find a working model."""
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except:
        return "models/gemini-1.5-flash"
    return "models/gemini-1.5-flash"

selected_model_name = find_working_model()
model = genai.GenerativeModel(selected_model_name)

def generate_pedati_plan(topic, syllabus, extra_context, links):
    """AI Prompt logic for PEDATI + Digital Citizenship + HOTS."""
    prompt = f"""
    Topic: {topic}. Syllabus Code: {syllabus}. 
    Tutor Resources/Links: {links}.
    Context: {extra_context}.

    Generate a comprehensive lesson plan in English. 
    Use these exact stage names: P [Prior Knowledge], E [Engage], D [Develop], A [Apply], T [Test], I [Improve].

    REQUIRED SECTIONS:
    SECTION: LESSON OBJECTIVES
    [4 points]
    SECTION: LESSON OUTCOMES
    [4 points]
    SECTION: SUCCESS CRITERIA
    [4 points]
    SECTION: PREREQUISITE
    [1 point]
    SECTION: KEYWORDS
    [10 items]
    
    SECTION: HOTS
    Generate 3 specific Higher Order Thinking Skills (HOTS) questions based on Bloom's Taxonomy.

    SECTION: DIGITAL CITIZENSHIP
    Based on {links}, generate 3-4 good digital habits.

    SECTION: PEDATI STAGES
    STAGE: P [Prior Knowledge] | SB: [Activity] | CB: [Activity]
    STAGE: E [Engage] | SB: [Activity] | CB: [Activity]
    STAGE: D [Develop] | SB: [Activity] | CB: [Activity]
    STAGE: A [Apply] | SB: [Activity] | CB: [Activity]
    STAGE: T [Test] | SB: [Activity] | CB: [Activity]
    STAGE: I [Improve] | SB: [Activity] | CB: [Activity]
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"System Error: {str(e)}"

def create_word_export(topic, syllabus, text, links):
    """Generates Word Doc with the 3-row stacked PEDATI format."""
    doc = Document()
    doc.add_heading(f'Lesson Plan: {topic} ({syllabus})', 0)

    # 1. Admin Header Table
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Week No :", "Date:"], ["No. of Students:", "Day:"], ["Venue / Lab No:", "Duration (mins):"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
    doc.add_paragraph()

    # 2. Resources Table
    doc.add_heading("Resources & Online Links", level=1)
    res_table = doc.add_table(rows=1, cols=1)
    res_table.style = 'Table Grid'
    res_table.cell(0, 0).text = f"Hardware: Smart board, Chromebook, Projector.\nLinks/Resources: {links}"

    # 3. Content Parsing
    sections = text.split('SECTION:')
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        title = lines[0].strip().replace("*", "").replace("#", "")
        content_lines = lines[1:]
        
        doc.add_heading(title.title(), level=1)
        
        # --- THE 3-ROW STACKED BOXING FOR PEDATI ---
        if "|" in section and "PEDATI" in title.upper():
            for line in content_lines:
                if "|" in line:
                    p = line.split("|")
                    
                    # Clean the text strings
                    stg = p[0].split(":")[-1].replace("*", "").replace("#", "").strip()
                    sb = p[1].split(":")[-1].replace("*", "").replace("#", "").strip()
                    cb = p[2].split(":")[-1].replace("*", "").replace("#", "").strip()
                    
                    # Create 3-row, 1-column table
                    table = doc.add_table(rows=3, cols=1)
                    table.style = 'Table Grid'
                    
                    # Row 1: Stage Title
                    cell0 = table.cell(0, 0)
                    cell0.text = f"STAGE: {stg}"
                    cell0.paragraphs[0].runs[0].bold = True
                    
                    # Row 2: SB Activity
                    table.cell(1, 0).text = f"SB: {sb}"
                    
                    # Row 3: CB Activity
                    table.cell(2, 0).text = f"CB: {cb}"
                    
                    doc.add_paragraph() # Spacer between stages

        # BOXING for HOTS & Digital Citizenship
        elif "DIGITAL CITIZENSHIP" in title.upper() or "HOTS" in title.upper():
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            prefix = "🧠 HOTS Questions:" if "HOTS" in title.upper() else "💡 Digital Habits:"
            clean_text = "\n".join([l.strip().replace("#", "").replace("*", "") for l in content_lines if l.strip()])
            table.cell(0,0).text = f"{prefix}\n" + clean_text
        
        # STANDARD BOXING for everything else
        else:
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            clean_text = "\n".join([l.strip().replace("#", "").replace("*", "") for l in content_lines if l.strip()])
            table.cell(0, 0).text = clean_text

    # 4. HOD Approval Section
    doc.add_page_break()
    doc.add_heading("HOD Approval & Remarks", level=1)
    hod_table = doc.add_table(rows=3, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Remark"; hod_table.cell(0, 1).text = "Signature / Stamp"
    hod_table.rows[1].height = Pt(60)
    hod_table.cell(2, 0).text = "Date:"; hod_table.cell(2, 1).text = "Name:"

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 2. GUI SECTION ---
st.set_page_config(page_title="PEDATI Master v2.2", layout="wide")
st.title("🎓 PEDATI Lesson Plan Generator v2.2")
st.info(f"System connected via: {selected_model_name}")
st.markdown("---")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Lesson Topic:")
with c2: u_syllabus = st.text_input("Syllabus Code:")

u_links = st.text_area("Online Resource Links:", placeholder="Paste links here...")
u_extra = st.text_area("Extra Context (Optional):")

if st.button("🚀 GENERATE PEDATI PLAN"):
    if u_topic and u_syllabus:
        with st.spinner("AI is crafting your upgraded PEDATI plan..."):
            result = generate_pedati_plan(u_topic, u_syllabus, u_extra, u_links)
            st.session_state['pedati_out'] = result

if 'pedati_out' in st.session_state:
    st.divider()
    st.text_area("AI Preview", st.session_state['pedati_out'], height=300)
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['pedati_out'], u_links)
    st.download_button("📥 Download Upgraded Word (.docx)", doc_file, f"PEDATI_V2.2_{u_topic}.docx")

st.markdown("---") 
st.markdown(
    """
    <div style='text-align: center; color: grey; font-size: 0.8em;'>
        <p><b>Smart PEDATI Lesson Plan AI-Generator v2.2</b></p>
        <p>Developed & Conceptualized by: <b>Hajah Nurul Haziqah @ Hjh Hartini Hj Nordin</b></p>
        <p>© 2026 PTES Academic Innovation Computer Science</p>
    </div>
    """,
    unsafe_allow_html=True
)
